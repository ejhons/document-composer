import os
from abc import ABC, abstractmethod
from typing import Dict, Type, Any
import pandas as pd
from docx import Document
import pypdfium2 as pdfium
import base64
import requests

# =====================================================================
# 1. THE PORT (Abstract Base Class)
# =====================================================================
class BaseContentAdapter(ABC):
    """
    The Port contract. Every format-specific adapter must implement this interface
    to guarantee interchangeable processing execution.
    """
    @abstractmethod
    def convert(self, source_path: str, output_dir: str) -> Any:
        """Processes the external file and returns data or handles post-processing."""
        pass


# =====================================================================
# 2. INDEPENDENT ADAPTER COMPONENT UNITS
# =====================================================================
class ExcelToMarkdownAdapter(BaseContentAdapter):
    def convert(self, source_path: str, output_dir: str) -> str:
        if not os.path.exists(source_path):
            raise FileNotFoundError(f"Spreadsheet asset missing: {source_path}")
        df = pd.read_excel(source_path).dropna(how='all')
        print(df)
        return f"\n\n{df.to_markdown(index=False)}\n\n"

class PdfToImageMarkdownAdapter(BaseContentAdapter):
    def convert(self, source_path: str, output_dir: str) -> str:
        if not os.path.exists(source_path):
            raise FileNotFoundError(f"PDF asset missing: {source_path}")
        
        os.makedirs(output_dir, exist_ok=True)
        base_name = os.path.splitext(os.path.basename(source_path))[0]
        
        print(f"[Adapter - Native] Opening PDF document via embedded Pdfium: {base_name}")
        
        # Abre o documento PDF de forma 100% nativa em Python
        pdf = pdfium.PdfDocument(source_path)
        image_tags = []
        
        # Percorre as páginas usando o índice
        for idx in range(len(pdf)):
            page = pdf[idx]
            
            # Renderiza a página direto para um objeto de Imagem do Pillow (PIL)
            # scale=2 equivale a aproximadamente 144 DPI, excelente qualidade para Word
            pil_image = page.render(scale=2).to_pil()
            
            # Define os caminhos de saída das imagens locais
            image_name = f"{base_name}_page_{idx + 1}.png"
            target_image_path = os.path.join(output_dir, image_name)
            
            # Salva o arquivo em disco
            pil_image.save(target_image_path, "PNG")
            
            # Formata as tags de injeção visual para o compilador Pandoc
            image_tags.append(f"![{base_name} - Page {idx + 1}]({target_image_path})\n\n---\n")
            
        return "\n\n" + "\n".join(image_tags) + "\n"

class MarkdownPassthroughAdapter(BaseContentAdapter):
    def convert(self, source_path: str, output_dir: str) -> str:
        if not os.path.exists(source_path):
            raise FileNotFoundError(f"External Markdown file missing: {source_path}")
        with open(source_path, 'r', encoding='utf-8') as file:
            return f"\n\n{file.read()}\n\n"


class DocxPostCompileAdapter(BaseContentAdapter):
    def convert(self, source_path: str, output_dir: str) -> str:
        # Returns a standard placement string token to be matched post-compilation
        return f"\n<!-- WORD_MERGE_POINT: {source_path} -->\n"

    def execute_binary_merge(self, master_docx_path: str, external_docx_path: str):
        """Specialized secondary action for docx post-processing layer."""
        if not os.path.exists(master_docx_path) or not os.path.exists(external_docx_path):
            raise FileNotFoundError("Missing target binaries for structural append.")
        
        master_doc = Document(master_docx_path)
        external_doc = Document(external_docx_path)
        master_doc.add_page_break()
        
        for element in external_doc.element.body:
            master_doc.element.body.append(element)
            
        master_doc.save(master_docx_path)

class MermaidMarkdownAdapter(BaseContentAdapter):
    """
    Specialist adapter that intercepts Mermaid code blocks, encodes them,
    and returns a standard Markdown image injection string.
    """
    def convert(self, source_path: str, output_dir: str) -> str:
        # Esse método cumpre o contrato da Porta abstrata se necessário, 
        # mas criaremos uma função específica para processar texto em lote (inline).
        pass

    @staticmethod
    def process_inline_diagram(mermaid_code: str, output_image_path: str) -> str:
        """
        Takes raw Mermaid text, sends it to the Mermaid.ink cloud rendering API,
        downloads the generated PNG, and returns the Markdown image syntax.
        """
        try:
            print("[Adapter] Generating vector graphic via Mermaid API...")
            
            # 1. Clean the code text
            clean_code = mermaid_code.strip()
            
            # 2. Convert text to standard base64 for URL transmission
            code_bytes = clean_code.encode('utf-8')
            base64_bytes = base64.b64encode(code_bytes)
            base64_string = base64_bytes.decode('utf-8')
            
            # 3. Request image from official cloud service
            # api_url = f"[https://mermaid.ink/img/](https://mermaid.ink/img/){base64_string}"
            api_url = f"https://mermaid.ink/img/{base64_string}"
            print(api_url)
            response = requests.get(api_url, timeout=10)
            
            if response.status_code == 200:
                # Save the image disk asset locally so Pandoc can read it offline
                os.makedirs(os.path.dirname(output_image_path), exist_ok=True)
                with open(output_image_path, 'wb') as img_file:
                    img_file.write(response.content)
                
                # Return standard Markdown image tag pointing to the local file
                return f"\n\n![System Diagram]({output_image_path})\n\n"
            else:
                print(f"[Warning] Mermaid API failed (Status: {response.status_code}). Falling back to code block.")
                return f"\n```mermaid\n{mermaid_code}\n```\n"
                
        except Exception as error:
            print(f"[Warning] Failed to render Mermaid diagram: {error}. Falling back to text.")
            return f"\n```mermaid\n{mermaid_code}\n```\n"

# =====================================================================
# 3. THE ADAPTER REGISTRY (The Registry Administrator)
# =====================================================================
class AdapterRegistry:
    """
    Manages the format-to-adapter mapping bindings dynamically.
    Acts as the single point of entry for format resolution.
    """
    def __init__(self):
        self._registry: Dict[str, BaseContentAdapter] = {}

    def register_adapter(self, format_extension: str, adapter: BaseContentAdapter):
        """Binds a specific file extension format to an adapter class instance."""
        self._registry[format_extension.lower()] = adapter

    def get_adapter(self, format_extension: str) -> BaseContentAdapter:
        """Retrieves the designated adapter instance or raises clean routing exceptions."""
        adapter = self._registry.get(format_extension.lower())
        if not adapter:
            raise ValueError(f"No adapter registered for format type: '{format_extension}'")
        return adapter